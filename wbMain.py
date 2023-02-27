from doctest import master
from tkinter import *
from tkinter import ttk
import tkinter.messagebox
from tkinter.tix import COLUMN
from turtle import width
import customtkinter
import webbrowser
import random
import time
from staticticForm import StaticticForm
import docx

import MySQLAdapt
from datetime import datetime
from docx import document

customtkinter.set_appearance_mode("Dark")
customtkinter.set_default_color_theme("green")


class App(customtkinter.CTk):
    WIDTH = 700
    HEIGHT = 590

    def __init__(self):
        super().__init__()

        self.title("ddex")
        self.geometry(f"{App.WIDTH}x{App.HEIGHT}")
     
        self.total_cost=0

        self.frame_left = customtkinter.CTkFrame(master=self, corner_radius=6)
        self.frame_left.grid(row=0, column=0, padx=10, pady=10)
        self.tovarrs = []

        customtkinter.CTkLabel(master=self.frame_left, text="    Клиент    ", bg_color='green').grid(row=0, column=0,
                                                                                                     columnspan=2,
                                                                                                     padx=10, pady=10,
                                                                                                     sticky="nwe")
        self.combobox1 = customtkinter.CTkComboBox(master=self.frame_left, values=[" "])
        self.combobox1.grid(row=0, column=2, columnspan=2, pady=10, padx=10)

        customtkinter.CTkButton(master=self.frame_left, text="+", command=lambda: AddClient(self)).grid(row=0, column=4,
                                                                                                        padx=10,
                                                                                                        pady=10,
                                                                                                        sticky="ew")

        customtkinter.CTkLabel(master=self.frame_left, text="    Дата    ", bg_color='green').grid(row=1, column=0,
                                                                                                   padx=10, pady=10,
                                                                                                   sticky="nwe")
        customtkinter.CTkLabel(master=self.frame_left, text=f"{datetime.now():%d-%m-%Y}", bg_color='green').grid(row=1,
                                                                                                                 column=2,
                                                                                                                 padx=10,
                                                                                                                 pady=10,
                                                                                                                 sticky="nwe")

        customtkinter.CTkLabel(master=self.frame_left, text="    Товар:    ", bg_color='green').grid(row=2, column=0,
                                                                                                     columnspan=2,
                                                                                                     padx=10, pady=10,
                                                                                                     sticky="nwe")
        self.combobox2 = customtkinter.CTkComboBox(master=self.frame_left, values=[" "])
        self.combobox2.grid(row=2, column=2, columnspan=2, pady=10, padx=10)
        FillComboBoxes(self)
        customtkinter.CTkButton(master=self.frame_left, text="+", command=lambda: AddTovar(self)).grid(row=2, column=4,
                                                                                                       padx=10, pady=10,
                                                                                                       sticky="ew")


        customtkinter.CTkLabel(master=self.frame_left, text="Количество товара:", bg_color='green').grid(row=4,
                                                                                                         column=0,
                                                                                                         columnspan=2,
                                                                                                         padx=10,
                                                                                                         pady=10,
                                                                                                         sticky="nwe")
        self.NumericUpDown = tkinter.Spinbox(master=self.frame_left, from_=0, to=999)
        self.NumericUpDown.grid(row=4, column=2, padx=10, pady=10, sticky="nwe")
        customtkinter.CTkButton(master=self.frame_left, text="В чек", command=lambda: AddToCheck(self)).grid(row=4,
                                                                                                               column=3,
                                                                                                               padx=10,
                                                                                                               pady=10,
                                                                                                               sticky="ew")

        customtkinter.CTkButton(master=self.frame_left, text="Чек", command=lambda: ShowCheck(self)).grid(row=6,
                                                                                                           column=0,
                                                                                                           padx=10,
                                                                                                           pady=10,
                                                                                                           sticky="ew")
        customtkinter.CTkButton(master=self.frame_left, text="Статистика", command=lambda: Stats(self)).grid(row=7,
                                                                                                             column=0,
                                                                                                             padx=10,
                                                                                                             pady=10,
                                                                                                             sticky="ew")
        self.listbox = tkinter.Listbox(master=self.frame_left)
        FillListbox(self)
        self.listbox.bind("<Return>", lambda x: delete_from_listbox(self))



def delete_from_listbox(self, *args):
    full_info = self.listbox.get(self.listbox.curselection()[0])
    id = full_info.split(":")[1].split(" ")[0]
    MySQLAdapt.SQL.delete_good(id)
    print(id)
    FillListbox(self)
    FillComboBoxes(self)


def AddClient(self):
    AddCli = customtkinter.CTkToplevel(self)
    AddCli.geometry("540x410")
    AddCli.title("ddex ADM")
    AddCli.frame_Q = customtkinter.CTkFrame(master=AddCli, width=450, height=450, border_width=4)
    AddCli.frame_Q.grid(row=22, column=6, sticky="nswe", padx=10, pady=10)

    customtkinter.CTkLabel(master=AddCli.frame_Q, text=" ").grid(row=0, column=0, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddCli.frame_Q, text=" ").grid(row=0, column=1, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddCli.frame_Q, text=" ").grid(row=0, column=2, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddCli.frame_Q, text=" ").grid(row=0, column=3, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddCli.frame_Q, text=" ").grid(row=0, column=4, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddCli.frame_Q, text=" ").grid(row=0, column=5, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddCli.frame_Q, text=" ").grid(row=0, column=6, padx=10, pady=10, sticky="nwe")

    customtkinter.CTkLabel(master=AddCli.frame_Q, text="Добавление нового клиента").grid(row=1, column=0, columnspan=3,
                                                                                         padx=10, pady=10, sticky="nwe")

    customtkinter.CTkLabel(master=AddCli.frame_Q, text="Фамилия - ").grid(row=2, column=0, padx=10, pady=10,
                                                                          sticky="nwe")
    AddCli.Fam = customtkinter.CTkEntry(master=AddCli.frame_Q, placeholder_text="Введите фамилию")
    AddCli.Fam.grid(row=2, column=1, columnspan=2, padx=10, pady=10, sticky="nwes")

    customtkinter.CTkLabel(master=AddCli.frame_Q, text="Имя - ").grid(row=3, column=0, padx=10, pady=10, sticky="nwe")
    AddCli.Name = customtkinter.CTkEntry(master=AddCli.frame_Q, placeholder_text="Введите имя")
    AddCli.Name.grid(row=3, column=1, columnspan=2, padx=10, pady=10, sticky="nwes")

    customtkinter.CTkLabel(master=AddCli.frame_Q, text="Отчество - ").grid(row=4, column=0, padx=10, pady=10,
                                                                           sticky="nwe")
    AddCli.Oth = customtkinter.CTkEntry(master=AddCli.frame_Q, placeholder_text="Введите отчество")
    AddCli.Oth.grid(row=4, column=1, columnspan=2, padx=10, pady=10, sticky="nwes")

    customtkinter.CTkLabel(master=AddCli.frame_Q, text="Дата - ").grid(row=5, column=0, padx=10, pady=10, sticky="nwe")
    AddCli.Data = customtkinter.CTkEntry(master=AddCli.frame_Q, placeholder_text="Введите дату в формате: дд.мм.гггг")
    AddCli.Data.grid(row=5, column=1, columnspan=5, padx=10, pady=10, sticky="nwes")

    customtkinter.CTkLabel(master=AddCli.frame_Q, text="Статус - ").grid(row=6, column=0, padx=10, pady=10,
                                                                         sticky="nwe")
    AddCli.Status = customtkinter.CTkComboBox(master=AddCli.frame_Q, values=["1", "0"])
    AddCli.Status.grid(row=6, column=1, pady=10, padx=10)

    ####################################################################             ,command=lambda: функция/метод юзающий БД пихать после закрывания кавычки "... в БД"
    customtkinter.CTkButton(master=AddCli.frame_Q, text="Добавить клиента в БД",
                            command=lambda: ADDDCLIENT(AddCli)).grid(row=7, column=0, columnspan=3, padx=10, pady=10,
                                                                     sticky="ew")



def ADDDCLIENT(AddCli):
    Fio = AddCli.Fam.get() + " " + AddCli.Name.get() + " " + AddCli.Oth.get()
    Dataa = AddCli.Data.get()
    Stat = AddCli.Status.get()
    MySQLAdapt.SQL.AddClient(Fio, Dataa, Stat)
    print(Fio)
    print(Dataa)
    print(Stat)
    AddCli.Fam.delete(0, tkinter.END)
    AddCli.Name.delete(0, tkinter.END)
    AddCli.Oth.delete(0, tkinter.END)
    AddCli.Data.delete(0, tkinter.END)
    AddCli.Stat.delete(0, tkinter.END)


def AddTovar(self):
    AddTov = customtkinter.CTkToplevel(self)
    AddTov.geometry("540x410")
    AddTov.title("ddex ADM")

    AddTov.frame_Z = customtkinter.CTkFrame(master=AddTov, width=450, height=450, border_width=4)
    AddTov.frame_Z.grid(row=22, column=6, sticky="nswe", padx=10, pady=10)

    customtkinter.CTkLabel(master=AddTov.frame_Z, text=" ").grid(row=0, column=0, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddTov.frame_Z, text=" ").grid(row=0, column=1, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddTov.frame_Z, text=" ").grid(row=0, column=2, padx=10, pady=10, sticky="nwe")
    customtkinter.CTkLabel(master=AddTov.frame_Z, text="Добавление нового товара").grid(row=1, column=0, columnspan=3,
                                                                                        padx=10, pady=10, sticky="nwe")

    customtkinter.CTkLabel(master=AddTov.frame_Z, text="Название - ").grid(row=2, column=0, padx=10, pady=10,
                                                                           sticky="nwe")
    AddTov.Name = customtkinter.CTkEntry(master=AddTov.frame_Z, placeholder_text="Введите название")
    AddTov.Name.grid(row=2, column=1, columnspan=2, padx=10, pady=10, sticky="nwes")

    customtkinter.CTkLabel(master=AddTov.frame_Z, text="Цена - ").grid(row=3, column=0, padx=10, pady=10, sticky="nwe")
    AddTov.Coast = customtkinter.CTkEntry(master=AddTov.frame_Z, placeholder_text="Введите цену")
    AddTov.Coast.grid(row=3, column=1, columnspan=2, padx=10, pady=10, sticky="nwes")

    customtkinter.CTkLabel(master=AddTov.frame_Z, text="Номер - ").grid(row=4, column=0, padx=10, pady=10, sticky="nwe")
    AddTov.Art = customtkinter.CTkEntry(master=AddTov.frame_Z, placeholder_text="Введите номер")
    AddTov.Art.grid(row=4, column=1, columnspan=2, padx=10, pady=10, sticky="nwes")

    customtkinter.CTkLabel(master=AddTov.frame_Z, text="Описание - ").grid(row=5, column=0, padx=10, pady=10,
                                                                           sticky="nwe")
    AddTov.Descr = customtkinter.CTkEntry(master=AddTov.frame_Z, placeholder_text="")
    AddTov.Descr.grid(row=5, column=1, columnspan=2, padx=10, pady=10, sticky="nwes")

    ####################################################################             ,command=lambda: функция/метод юзающий БД пихать после закрывания кавычки "... в БД"
    customtkinter.CTkButton(master=AddTov.frame_Z, text="Добавить товар в БД",
                            command=lambda: ADDBDTOVAR(AddTov, self)).grid(row=7, column=0, columnspan=3, padx=10,
                                                                           pady=10, sticky="ew")


def ADDBDTOVAR(AddTov, self):
    Name = AddTov.Name.get()
    Coast = AddTov.Coast.get()
    Art = AddTov.Art.get()
    Desc = AddTov.Descr.get()
    MySQLAdapt.SQL.AddProduct(Name, Coast, Art, Desc)
    print(Name)
    print(Coast)
    print(Art)
    print(Desc)
    FillListbox(self)




def AddToCheck(self):
      # Получаем индекс выбранной строки в Listbox
    selected_index = self.listbox.curselection()
    
    # Если строка не выбрана, выходим из функции
    if not selected_index:
        return
    
    # Получаем текст выбранной строки в Listbox
    selected_text = self.listbox.get(selected_index)

    
 


    # Получаем цену и номер(артикул) выбранного товара из текста выбранной строки
    price = float(selected_text.split("Цена:")[1].split(" руб.")[0])
    article_number = selected_text.split("Номер(артикул):")[1].split(" ")[0]
    
    # Получаем количество товара, выбранное в NumericUpDown
    quantity = int(self.NumericUpDown.get())
    
    # Вычисляем сумму товара
    total = price * quantity

    self.total_cost += total
    
    # Формируем строку с информацией о добавленном товаре и его стоимости
    check_item = f"Артикул:{article_number} Кол-во:{quantity} Сумма:{total} руб."
    
    # Добавляем строку в массив товаров на чеке
    self.tovarrs.append(check_item)


    tkinter.messagebox.showinfo("Готово", "Товар добавлен в чек.")

    
    
   
     
    

def ShowCheck(self):
   
    document = docx.Document()

# Добавляем заголовок
    document.add_heading('Чек', 0)

# Добавляем текущую дату и время
   
    
    document.add_paragraph('Дата: ' + str(f"{datetime.now():%d-%m-%Y}"))
    document.add_paragraph('Покупки: ' + str(self.tovarrs))
    document.add_paragraph('Клиент: ' + str(self.combobox1.get()))
    document.add_paragraph('К оплате: ' + str(self.total_cost) + "₽")

    document.save("test.docx")

    tkinter.messagebox.showinfo("Готово", "Чек создан.")
    self.total_cost = 0




def Stats(self):
    StaticticForm()


def FillListbox(self):
    self.listbox.delete(0, tkinter.END)
    # Получаем данные из таблицы product
    products = MySQLAdapt.SQL.GetALLTovar()

    # Создаем список строк, содержащих данные каждого продукта
    values = []
    for product in products:
        values.append(
            f"ID:{product[0]} Название:{product[1]} Цена:{product[2]} руб. Номер(артикул):{product[3]} Описание:{product[4]}")
        self.listbox.insert(tkinter.END, values[-1])


    # Создаем объект Variable, содержащий список строк
    values_var = tkinter.Variable(value=values)

    # Создаем объект Listbox, используя объект Variable
    self.listbox.grid(row=3, column=0, columnspan=5, padx=10, pady=10, sticky="ew")

    self.update()


def FillComboBoxes(self):
    ToComboBoxOne = MySQLAdapt.SQL.GetUser()
    values = list(map(lambda item: item[0], ToComboBoxOne))
    self.combobox1.configure(values=values)
    ToComboBoxTwo = MySQLAdapt.SQL.GetTovar()
    val = list(map(lambda item: item[0], ToComboBoxTwo))
    self.combobox2.configure(values=val)
    self.update()


if __name__ == "__main__":
    app = App()
    app.mainloop()
